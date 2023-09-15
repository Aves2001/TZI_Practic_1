import sys
import docx


def get_language(char: str):
    languages = {
        'Кирилиця': 'абвгдеєжзиіїйклмнопрстуфхцчшщьяёъыэ',
        'Латинська': 'abcdefghijklmnopqrstuvwxyz'
    }
    char_lower = char.lower()
    for language, alphabet in languages.items():
        if char_lower in alphabet:
            return alphabet
    return ''


def encrypt(text, shift):
    result = []
    for char in text:
        if char.isalpha():
            char_lower = char.lower()
            alphabet = get_language(char_lower)

            if char_lower in alphabet:
                index = (alphabet.index(char_lower) + shift) % len(alphabet)
                encrypted_char = alphabet[index]
                result.append(encrypted_char.upper() if char.isupper() else encrypted_char)
            else:
                result.append(char)
        else:
            result.append(char)
    return ''.join(result)


def get_key_caesar():
    while True:
        try:
            encryption_key = int(input("Введіть ключ шифрувння: "))
            return encryption_key
        except Exception:
            print("Введіть число")


def decrypt(text, key_caesar):
    return encrypt(text, -key_caesar)


def create_and_save_docx(file_path, content):
    doc = docx.Document()
    doc.add_paragraph(content)
    doc.save(file_path)


def save_text_to_file(text, output_suffix):
    if len(sys.argv) > 1:
        input_file_name = sys.argv[1]
        file_name, suffix = input_file_name.split('.')

        output_file_name = f"{file_name}{output_suffix}.{suffix}"
        if suffix == 'txt':
            with open(output_file_name, 'w', encoding='utf-8') as output_file:
                output_file.write(text)

        elif suffix == 'doc' or suffix == 'docx':
            create_and_save_docx(output_file_name, text)


def save_encrypted(encrypted_text):
    save_text_to_file(encrypted_text, '_encrypt')


def save_decrypted(decrypted_text):
    save_text_to_file(decrypted_text, '_decrypted')


def encrypt_or_decrypt(text, key_caesar):
    while True:
        print("1. Зашифрувти іформацію")
        print("2. Розшифрувти іформацію")
        action = input("Виберіть що потрібно зробити: ")

        print("Результут:")

        if action == '1':
            encrypted = encrypt(text, key_caesar)
            save_encrypted(encrypted)
            print(encrypted)
            break
        elif action == '2':
            decrypted = decrypt(text, key_caesar)
            save_decrypted(decrypted)
            print(decrypted)
            break


def read_docx(file_path):
    doc = docx.Document(file_path)
    text = ""
    for paragraph in doc.paragraphs:
        text += paragraph.text + "\n"
    return text


def main():
    if len(sys.argv) > 1:
        filename = sys.argv[1]
        suffix = filename.split('.')[1]
        if suffix == 'txt':
            with open(filename, encoding='utf-8') as f:
                text = f.read()
        elif suffix == 'doc' or suffix == 'docx':
            text = read_docx(filename)
    else:
        text = input("Введіть текст: ")

    key_caesar = get_key_caesar()
    encrypt_or_decrypt(text, key_caesar)

    _ = input()


if __name__ == "__main__":
    main()
